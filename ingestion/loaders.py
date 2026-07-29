"""
Document Loaders — Strategy & Factory patterns for loading multi-format documents.

Changes from v1
---------------
* ``CompositeDocumentLoader`` now accepts an optional ``PipelineConfig`` and
  uses a ``ThreadPoolExecutor`` inside ``stream_directory()`` to load files
  in parallel (I/O-bound work).
* ``load_directory()`` is preserved exactly as before — it now simply
  materialises the generator returned by ``stream_directory()``.
* Individual loaders (``PDFDocumentLoader``, etc.) are unchanged.
"""

from abc import ABC, abstractmethod
from concurrent.futures import ThreadPoolExecutor, as_completed
import csv
import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Iterator, List, Optional

try:
    import chardet
except ImportError:
    chardet = None
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from domain.models import Document
from ingestion.config import PipelineConfig
from ingestion.metrics import PipelineMetrics, StageTimer

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Abstract base
# ---------------------------------------------------------------------------

class BaseDocumentLoader(ABC):
    """Abstract Base Class for file format loaders."""

    @abstractmethod
    def can_load(self, file_extension: str) -> bool:
        """Return True if this loader handles the given extension."""
        pass

    @abstractmethod
    def load(self, filepath: str) -> List[Document]:
        """Extract text content and metadata from file into Document objects."""
        pass

    @staticmethod
    def _get_modified_date(filepath: str) -> str:
        stat = os.stat(filepath)
        return datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat()


# ---------------------------------------------------------------------------
# Concrete loaders (unchanged from v1)
# ---------------------------------------------------------------------------

class PDFDocumentLoader(BaseDocumentLoader):
    """Loader for PDF documents using PyMuPDF."""

    def can_load(self, file_extension: str) -> bool:
        return file_extension.lower() == ".pdf"

    def load(self, filepath: str) -> List[Document]:
        if fitz is None:
            logger.error("PyMuPDF (fitz) is not installed.")
            return []
        documents = []
        try:
            doc = fitz.open(filepath)
            modified_date = self._get_modified_date(filepath)
            file_name = os.path.basename(filepath)
            abs_path = os.path.abspath(filepath)

            for page_num, page in enumerate(doc, start=1):
                text = page.get_text("text").strip()
                if text:
                    documents.append(
                        Document(
                            text=text,
                            source_file=abs_path,
                            file_name=file_name,
                            file_type=".pdf",
                            page_number=page_num,
                            total_pages=len(doc),
                            modified_date=modified_date,
                        )
                    )
            doc.close()
        except Exception as e:
            logger.error("Failed to read PDF '%s': %s", filepath, e)
        return documents


class DocxDocumentLoader(BaseDocumentLoader):
    """Loader for Microsoft Word (.docx) files."""

    def can_load(self, file_extension: str) -> bool:
        return file_extension.lower() == ".docx"

    def load(self, filepath: str) -> List[Document]:
        if DocxDocument is None:
            logger.error("python-docx is not installed.")
            return []
        try:
            doc = DocxDocument(filepath)
            full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            if full_text.strip():
                return [
                    Document(
                        text=full_text,
                        source_file=os.path.abspath(filepath),
                        file_name=os.path.basename(filepath),
                        file_type=".docx",
                        modified_date=self._get_modified_date(filepath),
                    )
                ]
        except Exception as e:
            logger.error("Failed to read DOCX '%s': %s", filepath, e)
        return []


class TextDocumentLoader(BaseDocumentLoader):
    """Loader for plain text (.txt) and Markdown (.md) files."""

    def can_load(self, file_extension: str) -> bool:
        return file_extension.lower() in {".txt", ".md"}

    def load(self, filepath: str) -> List[Document]:
        try:
            with open(filepath, "rb") as f:
                raw = f.read()
            encoding = "utf-8"
            if chardet is not None:
                encoding = chardet.detect(raw).get("encoding", "utf-8") or "utf-8"
            text = raw.decode(encoding, errors="replace").strip()

            if text:
                ext = Path(filepath).suffix.lower()
                return [
                    Document(
                        text=text,
                        source_file=os.path.abspath(filepath),
                        file_name=os.path.basename(filepath),
                        file_type=ext,
                        modified_date=self._get_modified_date(filepath),
                    )
                ]
        except Exception as e:
            logger.error("Failed to read text file '%s': %s", filepath, e)
        return []


class CSVDocumentLoader(BaseDocumentLoader):
    """Loader for CSV documents."""

    def can_load(self, file_extension: str) -> bool:
        return file_extension.lower() == ".csv"

    def load(self, filepath: str) -> List[Document]:
        try:
            with open(filepath, "rb") as f:
                raw = f.read()
            encoding = "utf-8"
            if chardet is not None:
                encoding = chardet.detect(raw).get("encoding", "utf-8") or "utf-8"
            text = raw.decode(encoding, errors="replace")

            reader = csv.DictReader(text.splitlines())
            rows = []
            for row in reader:
                row_str = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                if row_str.strip():
                    rows.append(row_str)

            if rows:
                return [
                    Document(
                        text="\n".join(rows),
                        source_file=os.path.abspath(filepath),
                        file_name=os.path.basename(filepath),
                        file_type=".csv",
                        modified_date=self._get_modified_date(filepath),
                    )
                ]
        except Exception as e:
            logger.error("Failed to read CSV '%s': %s", filepath, e)
        return []


# ---------------------------------------------------------------------------
# Composite loader
# ---------------------------------------------------------------------------

class CompositeDocumentLoader:
    """Composite Loader orchestrating specialised document loaders.

    Parameters
    ----------
    loaders:
        Ordered list of ``BaseDocumentLoader`` instances.  The first loader
        that returns ``can_load=True`` for a file extension is used.
    config:
        ``PipelineConfig`` controlling worker counts and metrics collection.
        Defaults to a ``PipelineConfig()`` with library defaults if omitted.
    """

    def __init__(
        self,
        loaders: Optional[List[BaseDocumentLoader]] = None,
        config: Optional[PipelineConfig] = None,
    ) -> None:
        self.loaders = loaders or [
            PDFDocumentLoader(),
            DocxDocumentLoader(),
            TextDocumentLoader(),
            CSVDocumentLoader(),
        ]
        self.config = config or PipelineConfig()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def load_directory(self, folder_path: str) -> List[Document]:
        """Scan *folder_path* recursively and return all supported documents.

        **Backward-compatible** — preserves the original batch API.
        Internally delegates to :meth:`stream_directory` and materialises
        the result into a list.
        """
        return list(self.stream_directory(folder_path))

    def stream_directory(
        self,
        folder_path: str,
        metrics: Optional[PipelineMetrics] = None,
        state_store=None,  # Optional[BaseCheckpointStore] — avoids circular import
    ) -> Iterator[Document]:
        """Yield :class:`~domain.models.Document` objects as they are loaded.

        Files are submitted to a ``ThreadPoolExecutor`` (size controlled by
        ``config.loader_workers``) and results are yielded as futures
        complete, so downstream stages can start processing before all files
        are read.

        Parameters
        ----------
        folder_path:
            Absolute or relative path to a directory.
        metrics:
            Optional :class:`~ingestion.metrics.PipelineMetrics` instance.
            ``files_found``, ``files_skipped``, ``docs_loaded``, and
            ``load_errors`` are updated in-place.
        state_store:
            Optional checkpoint store.  When provided, files whose SHA-256
            hash has not changed since the last successful run are skipped
            and ``metrics.files_skipped`` is incremented.

        Raises
        ------
        FileNotFoundError
            If *folder_path* does not exist or is not a directory.
        """
        folder = Path(folder_path)
        if not folder.is_dir():
            raise FileNotFoundError(f"Folder not found: {folder_path}")

        paths = self._discover_files(folder)
        if metrics:
            metrics.files_found = len(paths)

        with StageTimer(metrics or PipelineMetrics(), "load"):
            with ThreadPoolExecutor(max_workers=self.config.loader_workers) as pool:
                future_to_path = {}
                for path in paths:
                    # Incremental skip: delegate to state_store if provided
                    if state_store is not None and not state_store.is_file_changed(path):
                        if metrics:
                            metrics.files_skipped += 1
                        logger.debug("loader: skipping unchanged file %s", path)
                        continue
                    future_to_path[pool.submit(self._load_file, path)] = path

                for future in as_completed(future_to_path):
                    path = future_to_path[future]
                    try:
                        docs = future.result()
                        if metrics:
                            if docs:
                                metrics.docs_loaded += len(docs)
                            else:
                                metrics.load_errors += 1
                        yield from docs
                    except Exception as exc:
                        logger.error("loader: unhandled error for '%s': %s", path, exc)
                        if metrics:
                            metrics.load_errors += 1

        if metrics:
            logger.info(
                "stage=load files_found=%d skipped=%d docs=%d errors=%d time_s=%.2f",
                metrics.files_found,
                metrics.files_skipped,
                metrics.docs_loaded,
                metrics.load_errors,
                metrics.load_time_s,
            )

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _discover_files(self, folder: Path) -> List[str]:
        """Return a sorted list of absolute file paths under *folder*."""
        paths = []
        for root, _, filenames in os.walk(folder):
            for fname in filenames:
                ext = Path(fname).suffix.lower()
                if any(ldr.can_load(ext) for ldr in self.loaders):
                    paths.append(os.path.join(root, fname))
        return sorted(paths)

    def _load_file(self, filepath: str) -> List[Document]:
        """Load a single file using the appropriate loader."""
        ext = Path(filepath).suffix.lower()
        for loader in self.loaders:
            if loader.can_load(ext):
                return loader.load(filepath)
        return []
